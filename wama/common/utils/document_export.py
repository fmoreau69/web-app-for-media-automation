"""
WAMA — Document Export Utilities
Génération de fichiers PDF (fpdf2) et DOCX (python-docx) à partir des résultats
des applications Describer et Transcriber.
"""
import io
import re
import datetime


# ──────────────────────────────────────────────────────────────────────────────
# Helper — FPDF base class with header/footer
# ──────────────────────────────────────────────────────────────────────────────

def _strip_markdown(text: str) -> str:
    """Strip markdown syntax while preserving paragraph structure for PDF/DOCX rendering."""
    if not text:
        return text
    t = text
    # Remove heading markers
    t = re.sub(r'^#{1,6}\s+', '', t, flags=re.MULTILINE)
    # Remove bold/italic markers
    t = re.sub(r'\*{1,3}|_{1,3}', '', t)
    # Remove table separator lines (|---|---|)
    t = re.sub(r'^\s*\|[-|: ]+\|\s*$', '', t, flags=re.MULTILINE)
    # Replace table pipes with spaces
    t = re.sub(r'\|', '  ', t)
    # Remove inline code backticks
    t = re.sub(r'`+', '', t)
    # Strip list bullet markers at line start
    t = re.sub(r'^\s*[-*+]\s+', '', t, flags=re.MULTILINE)
    # Collapse 3+ blank lines into 2
    t = re.sub(r'\n{3,}', '\n\n', t)
    # Strip trailing whitespace per line
    t = re.sub(r'[ \t]+$', '', t, flags=re.MULTILINE)
    return t.strip()


def _sanitize_for_latin1(text: str) -> str:
    """Replace characters outside Latin-1 with ASCII equivalents for fpdf2 built-in fonts."""
    if not text:
        return text
    replacements = {
        '\u2014': '-', '\u2013': '-',        # em-dash, en-dash
        '\u2018': "'", '\u2019': "'",        # curly single quotes
        '\u201c': '"', '\u201d': '"',        # curly double quotes
        '\u2026': '...', '\u2022': '*',      # ellipsis, bullet
        '\u2192': '->', '\u2190': '<-',      # arrows
        '\u00b0': ' deg',                    # degree sign (already Latin-1 but safe)
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    # Encode to Latin-1, replacing any remaining non-Latin-1 chars with '?'
    return text.encode('latin-1', errors='replace').decode('latin-1')


def _make_pdf():
    """Return a configured FPDF2 instance with header/footer."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise ImportError("fpdf2 n'est pas installé. Exécutez : pip install fpdf2")

    class WamaPDF(FPDF):
        def header(self):
            self.set_font('Helvetica', 'B', 8)
            self.set_text_color(140, 140, 140)
            self.cell(0, 5, 'WAMA - Export', align='R')
            self.ln(3)

        def footer(self):
            self.set_y(-12)
            self.set_font('Helvetica', '', 8)
            self.set_text_color(160, 160, 160)
            self.cell(0, 8, f'Page {self.page_no()}', align='C')

    pdf = WamaPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(18, 18, 18)
    return pdf


def _section_title(pdf, text):
    """Print a section title line."""
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_text_color(40, 120, 200)
    pdf.cell(0, 6, text, ln=True)
    pdf.set_draw_color(40, 120, 200)
    pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + pdf.epw, pdf.get_y())
    pdf.set_draw_color(0, 0, 0)
    pdf.ln(3)
    pdf.set_text_color(30, 30, 30)


def _meta_line(pdf, label, value):
    """Print a key: value metadata line."""
    pdf.set_font('Helvetica', 'B', 9)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(35, 5, label + ' :', ln=False)
    pdf.set_font('Helvetica', '', 9)
    pdf.set_text_color(30, 30, 30)
    pdf.multi_cell(0, 5, str(value))


def _body_text(pdf, text, font_size=10):
    """Print multi-line body text with word-wrap."""
    pdf.set_font('Helvetica', '', font_size)
    pdf.set_text_color(30, 30, 30)
    # fpdf2 multi_cell handles Unicode and word-wrap
    pdf.multi_cell(0, 5, text or '')
    pdf.ln(2)


def _bullet_list(pdf, items):
    """Print a bulleted list from a Python list of strings."""
    pdf.set_font('Helvetica', '', 10)
    pdf.set_text_color(30, 30, 30)
    for item in items:
        pdf.cell(6, 5, '\u2022', ln=False)
        pdf.multi_cell(0, 5, str(item))
    pdf.ln(2)


# ──────────────────────────────────────────────────────────────────────────────
# DESCRIBER — PDF
# ──────────────────────────────────────────────────────────────────────────────

_FORMAT_TITLES = {
    'summary':       'RÉSUMÉ',
    'detailed':      'DESCRIPTION DÉTAILLÉE',
    'scientific':    'SYNTHÈSE SCIENTIFIQUE',
    'bullet_points': 'POINTS CLÉS',
    'meeting':       'COMPTE-RENDU DE RÉUNION',
}


def generate_description_pdf(description) -> bytes:
    """
    Generate a PDF from a Description instance.
    Adapts layout to description.output_format.
    Returns raw PDF bytes.
    """
    pdf = _make_pdf()
    pdf.add_page()

    doc_title = _FORMAT_TITLES.get(description.output_format, 'DESCRIPTION')

    # ── Title block ──
    pdf.set_font('Helvetica', 'B', 16)
    pdf.set_text_color(20, 20, 20)
    pdf.cell(0, 10, doc_title, align='C', ln=True)
    pdf.ln(2)

    # ── Metadata ──
    _section_title(pdf, 'Informations')
    _meta_line(pdf, 'Fichier source', description.filename or '—')
    _meta_line(pdf, 'Type', description.detected_type or description.content_type)
    _meta_line(pdf, 'Format sortie', doc_title.title())
    _meta_line(pdf, 'Langue', description.get_output_language_display())
    _meta_line(pdf, 'Date', description.created_at.strftime('%d/%m/%Y %H:%M'))
    pdf.ln(4)

    # ── Main result ──
    if description.output_format == 'bullet_points':
        _section_title(pdf, 'Points Clés')
        lines = [l.lstrip('•- ').strip() for l in (description.result_text or '').splitlines() if l.strip()]
        _bullet_list(pdf, lines)
    else:
        _section_title(pdf, doc_title.title())
        _body_text(pdf, description.result_text or '')

    # ── Optional summary ──
    if description.summary:
        pdf.ln(2)
        _section_title(pdf, 'Résumé')
        _body_text(pdf, description.summary)

    # ── Optional coherence ──
    if description.coherence_score is not None:
        pdf.ln(2)
        _section_title(pdf, f'Cohérence — Score {description.coherence_score}/10')
        if description.coherence_notes:
            _body_text(pdf, description.coherence_notes)
        if description.coherence_suggestion:
            pdf.set_font('Helvetica', 'I', 9)
            pdf.set_text_color(100, 100, 100)
            pdf.multi_cell(0, 5, 'Suggestion : ' + description.coherence_suggestion)

    return bytes(pdf.output())


# ──────────────────────────────────────────────────────────────────────────────
# DESCRIBER — DOCX
# ──────────────────────────────────────────────────────────────────────────────

def generate_description_docx(description) -> bytes:
    """Generate a DOCX from a Description instance. Returns raw bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx n'est pas installé. Exécutez : pip install python-docx")

    doc = Document()

    doc_title = _FORMAT_TITLES.get(description.output_format, 'DESCRIPTION')
    doc.add_heading(doc_title, 0)

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    rows_data = [
        ('Fichier source', description.filename or '—'),
        ('Type', description.detected_type or description.content_type),
        ('Langue', description.get_output_language_display()),
        ('Date', description.created_at.strftime('%d/%m/%Y %H:%M')),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_paragraph()

    # Main result
    doc.add_heading(doc_title.title(), 1)
    if description.output_format == 'bullet_points':
        for line in (description.result_text or '').splitlines():
            line = line.lstrip('•- ').strip()
            if line:
                doc.add_paragraph(line, style='List Bullet')
    else:
        doc.add_paragraph(description.result_text or '')

    # Optional summary
    if description.summary:
        doc.add_heading('Résumé', 1)
        doc.add_paragraph(description.summary)

    # Optional coherence
    if description.coherence_score is not None:
        doc.add_heading(f'Cohérence — {description.coherence_score}/10', 1)
        if description.coherence_notes:
            doc.add_paragraph(description.coherence_notes)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────────────────────────────────────
# TRANSCRIBER — PDF
# ──────────────────────────────────────────────────────────────────────────────

def _transcript_blocks(transcript):
    """Blocs de transcription compactés par locuteur (segments consécutifs d'un même
    locuteur fusionnés) pour un document de sortie lisible.

    Retourne une liste de dicts {speaker, start, end, text}. Les libellés sont déjà
    résolus via speaker_map (nom personnalisé si défini). Fallback : [] si pas de segments.
    """
    from wama.transcriber.models import TranscriptSegment
    from wama.transcriber.utils.speakers import display_speaker
    spk_map = transcript.speaker_map or {}
    segs = list(TranscriptSegment.objects.filter(transcript=transcript).order_by('order'))
    blocks = []
    for s in segs:
        spk = display_speaker(s.speaker_id, spk_map)
        txt = (s.text or '').strip()
        if not txt:
            continue
        if blocks and blocks[-1]['speaker'] == spk:
            # Même locuteur que le bloc précédent → on fusionne (compactage).
            blocks[-1]['text'] += ' ' + txt
            blocks[-1]['end'] = s.end_time
        else:
            blocks.append({'speaker': spk, 'start': s.start_time, 'end': s.end_time, 'text': txt})
    return blocks


def _transcript_speakers(transcript):
    """Liste ordonnée des intervenants (noms d'affichage) présents dans la transcription."""
    from wama.transcriber.models import TranscriptSegment
    from wama.transcriber.utils.speakers import unique_speakers, display_speaker
    spk_map = transcript.speaker_map or {}
    rows = TranscriptSegment.objects.filter(transcript=transcript).order_by('order')
    canon = unique_speakers([{'speaker_id': r.speaker_id} for r in rows])
    return [display_speaker(c, spk_map) for c in canon]


def generate_transcript_txt(transcript) -> bytes:
    """Génère un .txt mis en forme : avant-propos (titre/date/intervenants) + blocs
    compactés par locuteur. Fallback sur transcript.text brut si pas de segments."""
    lines = []
    title = (transcript.title or '').strip() or (getattr(transcript, 'filename', '') or 'Transcription')
    lines.append(title)
    lines.append('=' * len(title))
    if transcript.meeting_date:
        lines.append(f"Date : {transcript.meeting_date.strftime('%d/%m/%Y')}")
    stem = getattr(transcript, 'filename', None) or '—'
    lines.append(f"Fichier source : {stem}")
    if transcript.duration_display:
        lines.append(f"Durée : {transcript.duration_display}")
    speakers = _transcript_speakers(transcript)
    if speakers:
        lines.append("Intervenants : " + ', '.join(speakers))
    lines.append('')

    blocks = _transcript_blocks(transcript)
    if blocks:
        for b in blocks:
            ts = f"[{_fmt_time(b['start'])} — {_fmt_time(b['end'])}]"
            lines.append(f"{b['speaker'] or 'SPK'} {ts}")
            lines.append(b['text'])
            lines.append('')
    else:
        lines.append(transcript.text or '')
    return ('\n'.join(lines)).encode('utf-8')


def generate_transcript_pdf(transcript) -> bytes:
    """Generate a PDF from a Transcript instance. Returns raw PDF bytes."""
    pdf = _make_pdf()
    pdf.add_page()

    # Detect if meeting report format
    is_meeting = (transcript.summary_type == 'meeting') if transcript.summary_type else False
    custom_title = (transcript.title or '').strip()
    doc_title = custom_title or ('COMPTE-RENDU DE RÉUNION' if is_meeting else 'TRANSCRIPTION')

    # ── Title ──
    pdf.set_font('Helvetica', 'B', 16)
    pdf.set_text_color(20, 20, 20)
    pdf.cell(0, 10, _sanitize_for_latin1(doc_title), align='C', ln=True)
    pdf.ln(2)

    # ── Metadata / avant-propos ──
    _section_title(pdf, 'Informations')
    stem = getattr(transcript, 'filename', None) or (
        transcript.audio.name.split('/')[-1] if transcript.audio else '—'
    )
    if transcript.meeting_date:
        _meta_line(pdf, 'Date', transcript.meeting_date.strftime('%d/%m/%Y'))
    _meta_line(pdf, 'Fichier source', stem)
    _meta_line(pdf, 'Durée', transcript.duration_display or '—')
    _meta_line(pdf, 'Backend', transcript.used_backend or transcript.backend or 'auto')
    _meta_line(pdf, 'Langue', transcript.language or 'auto')
    speakers = _transcript_speakers(transcript)
    if speakers:
        _meta_line(pdf, 'Intervenants', _sanitize_for_latin1(', '.join(speakers)))
    pdf.ln(4)

    # ── Summary / key points (if meeting) ──
    if is_meeting and transcript.summary:
        _section_title(pdf, 'Résumé')
        _body_text(pdf, transcript.summary)
        pdf.ln(2)

    if transcript.key_points:
        kp = transcript.key_points
        if isinstance(kp, list) and kp:
            _section_title(pdf, 'Points Clés')
            _bullet_list(pdf, kp)
            pdf.ln(2)

    if transcript.action_items:
        ai = transcript.action_items
        if isinstance(ai, list) and ai:
            _section_title(pdf, 'Actions')
            _bullet_list(pdf, ai)
            pdf.ln(2)

    # ── Full transcript (compacté par locuteur) ──
    _section_title(pdf, 'Transcription Complète')

    try:
        blocks = _transcript_blocks(transcript)
        if blocks:
            for b in blocks:
                pdf.set_font('Helvetica', 'B', 9)
                pdf.set_text_color(40, 120, 200)
                time_str = f"{_fmt_time(b['start'])} — {_fmt_time(b['end'])}"
                pdf.cell(0, 5, _sanitize_for_latin1(f"[{b['speaker'] or 'SPK'}]  {time_str}"), ln=True)
                pdf.set_font('Helvetica', '', 10)
                pdf.set_text_color(30, 30, 30)
                pdf.multi_cell(0, 5, b['text'])
                pdf.ln(1)
        else:
            _body_text(pdf, transcript.text or '')
    except Exception:
        _body_text(pdf, transcript.text or '')

    return bytes(pdf.output())


def _fmt_time(seconds):
    """Format seconds → MM:SS."""
    if seconds is None:
        return '??:??'
    m, s = divmod(int(seconds), 60)
    return f'{m:02d}:{s:02d}'


# ──────────────────────────────────────────────────────────────────────────────
# READER — PDF
# ──────────────────────────────────────────────────────────────────────────────

def generate_reader_pdf(item) -> bytes:
    """Generate a PDF from a ReadingItem instance. Returns raw PDF bytes."""
    pdf = _make_pdf()
    pdf.add_page()

    # ── Title ──
    pdf.set_font('Helvetica', 'B', 16)
    pdf.set_text_color(20, 20, 20)
    pdf.cell(0, 10, 'DOCUMENT OCR', align='C', ln=True)
    pdf.ln(2)

    # ── Metadata ──
    _section_title(pdf, 'Informations')
    _meta_line(pdf, 'Fichier source', _sanitize_for_latin1(item.filename or '-'))
    _meta_line(pdf, 'Pages', str(item.page_count) if item.page_count else '-')
    _meta_line(pdf, 'Backend', item.used_backend or item.backend or 'auto')
    _meta_line(pdf, 'Mode', item.mode or 'auto')
    if item.language:
        _meta_line(pdf, 'Langue', item.language)
    _meta_line(pdf, 'Date', item.created_at.strftime('%d/%m/%Y %H:%M'))
    pdf.ln(4)

    # ── OCR Text ──
    _section_title(pdf, 'Texte Extrait')
    _body_text(pdf, _sanitize_for_latin1(_strip_markdown(item.result_text or '')))

    # ── Analysis (if any) ──
    if item.analysis:
        pdf.ln(2)
        _section_title(pdf, 'Analyse')
        _body_text(pdf, _sanitize_for_latin1(_strip_markdown(item.analysis)))

    return bytes(pdf.output())


# ──────────────────────────────────────────────────────────────────────────────
# READER — DOCX
# ──────────────────────────────────────────────────────────────────────────────

def generate_reader_docx(item) -> bytes:
    """Generate a DOCX from a ReadingItem instance. Returns raw bytes."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx n'est pas installé. Exécutez : pip install python-docx")

    doc = Document()
    doc.add_heading('Document OCR', 0)

    # Metadata table
    rows_data = [
        ('Fichier source', item.filename or '—'),
        ('Pages', str(item.page_count) if item.page_count else '—'),
        ('Backend', item.used_backend or item.backend or 'auto'),
        ('Mode', item.mode or 'auto'),
        ('Date', item.created_at.strftime('%d/%m/%Y %H:%M')),
    ]
    if item.language:
        rows_data.insert(3, ('Langue', item.language))

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_paragraph()

    # OCR text
    doc.add_heading('Texte Extrait', 1)
    doc.add_paragraph(_strip_markdown(item.result_text or ''))

    # Analysis
    if item.analysis:
        doc.add_heading('Analyse', 1)
        doc.add_paragraph(_strip_markdown(item.analysis))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────────────────────────────────────
# TRANSCRIBER — DOCX
# ──────────────────────────────────────────────────────────────────────────────

def generate_transcript_docx(transcript) -> bytes:
    """Generate a DOCX from a Transcript instance. Returns raw bytes."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx n'est pas installé. Exécutez : pip install python-docx")

    doc = Document()
    is_meeting = (transcript.summary_type == 'meeting') if transcript.summary_type else False
    custom_title = (transcript.title or '').strip()
    doc_title = custom_title or ('Compte-Rendu de Réunion' if is_meeting else 'Transcription')
    doc.add_heading(doc_title, 0)

    stem = getattr(transcript, 'filename', None) or (
        transcript.audio.name.split('/')[-1] if transcript.audio else '—'
    )

    # Metadata table / avant-propos
    rows_data = []
    if transcript.meeting_date:
        rows_data.append(('Date', transcript.meeting_date.strftime('%d/%m/%Y')))
    rows_data += [
        ('Fichier source', stem),
        ('Durée', transcript.duration_display or '—'),
        ('Backend', transcript.used_backend or 'auto'),
    ]
    speakers = _transcript_speakers(transcript)
    if speakers:
        rows_data.append(('Intervenants', ', '.join(speakers)))
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_paragraph()

    # Summary
    if is_meeting and transcript.summary:
        doc.add_heading('Résumé', 1)
        doc.add_paragraph(transcript.summary)

    # Key points
    if transcript.key_points and isinstance(transcript.key_points, list):
        doc.add_heading('Points Clés', 1)
        for kp in transcript.key_points:
            doc.add_paragraph(str(kp), style='List Bullet')

    # Action items
    if transcript.action_items and isinstance(transcript.action_items, list):
        doc.add_heading('Actions', 1)
        for ai in transcript.action_items:
            doc.add_paragraph(str(ai), style='List Bullet')

    # Full transcript (compacté par locuteur)
    doc.add_heading('Transcription Complète', 1)
    try:
        blocks = _transcript_blocks(transcript)
        if blocks:
            for b in blocks:
                p = doc.add_paragraph()
                run = p.add_run(f"[{b['speaker'] or 'SPK'}]  {_fmt_time(b['start'])}—{_fmt_time(b['end'])}\n")
                run.bold = True
                p.add_run(b['text'])
        else:
            doc.add_paragraph(transcript.text or '')
    except Exception:
        doc.add_paragraph(transcript.text or '')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
